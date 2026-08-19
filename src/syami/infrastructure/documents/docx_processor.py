from pathlib import Path

from docx import Document

from syami.application.documents.processor import DocumentProcessor
from syami.domain.document import ContentUnit, ProcessedDocument


class DOCXProcessor(DocumentProcessor):

    def process(
        self,
        document_id: int,
        source_path: str,
    ) -> ProcessedDocument:

        path = Path(source_path)

        if not path.exists():
            raise FileNotFoundError(
                f"File does not exist: {path}"
            )

        if path.suffix.lower() != ".docx":
            raise ValueError(
                f"Expected a DOCX file: {path}"
            )

        document = Document(str(path))

        units = []

        for index, paragraph in enumerate(
            document.paragraphs,
            start=1,
        ):
            text = paragraph.text.strip()

            if not text:
                continue

            units.append(
                ContentUnit(
                    text=text,
                    unit_type="paragraph",
                    unit_index=index,
                )
            )

        return ProcessedDocument(
            document_id=document_id,
            source_path=str(path),
            document_type="docx",
            title=path.stem,
            units=units,
        )